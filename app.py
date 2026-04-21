import os
import re
import json
import base64
import tempfile
from difflib import SequenceMatcher
from functools import lru_cache
from collections import defaultdict
from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS
import pdfplumber
import docx
from bs4 import BeautifulSoup
import io
from PIL import Image, ImageOps, ImageFilter
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

app = Flask(__name__, template_folder='.')
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max

# ──────────────────────────────────────────────
#  SDG Definitions & Keywords
# ──────────────────────────────────────────────
SDG_INFO = {
    1:  {"name": "No Poverty",                              "color": "#E5243B"},
    2:  {"name": "Zero Hunger",                             "color": "#DDA63A"},
    3:  {"name": "Good Health and Well-Being",              "color": "#4C9F38"},
    4:  {"name": "Quality Education",                       "color": "#C5192D"},
    5:  {"name": "Gender Equality",                         "color": "#FF3A21"},
    6:  {"name": "Clean Water and Sanitation",              "color": "#26BDE2"},
    7:  {"name": "Affordable and Clean Energy",             "color": "#FCC30B"},
    8:  {"name": "Decent Work and Economic Growth",         "color": "#A21942"},
    9:  {"name": "Industry, Innovation and Infrastructure", "color": "#FD6925"},
    10: {"name": "Reduced Inequalities",                    "color": "#DD1367"},
    11: {"name": "Sustainable Cities and Communities",      "color": "#FD9D24"},
    12: {"name": "Responsible Consumption and Production",  "color": "#BF8B2E"},
    13: {"name": "Climate Action",                          "color": "#3F7E44"},
    14: {"name": "Life Below Water",                        "color": "#0A97D9"},
    15: {"name": "Life on Land",                            "color": "#56C02B"},
    16: {"name": "Peace, Justice and Strong Institutions",  "color": "#00689D"},
    17: {"name": "Partnerships for the Goals",              "color": "#19486A"},
}

SDG_KEYWORDS = {
    1:  ["poverty", "poor", "income", "welfare", "social protection", "inequality", "deprivation",
         "livelihood", "subsistence", "destitute", "homeless", "underprivileged", "financial inclusion",
         "basic needs", "minimum wage", "safety net", "food stamps", "cash transfer", "poverty alleviation",
         "rural development", "microfinance", "financial literacy", "economic empowerment", "inclusive growth",
         "welfare policies", "community development", "employment generation", "resource equity",
         "sustainable livelihoods", "social entrepreneurship", "economic resilience", "eradication strategies"],
    2:  ["hunger", "food security", "malnutrition", "famine", "agriculture", "crop", "farming",
         "nutrition", "starvation", "food system", "agroforestry", "smallholder", "food production",
         "dietary", "fertilizer", "irrigation", "seed", "harvest", "food waste", "aquaculture",
         "sustainable agriculture", "agro ecology", "crop diversification", "soil fertility",
         "food processing", "agri business", "food distribution", "climate smart farming",
         "post harvest technology", "organic farming", "farmer empowerment", "precision agriculture",
         "animal husbandry", "cold storage plants"],
    3:  ["health", "healthcare", "disease", "medicine", "hospital", "vaccination", "mental health",
         "mortality", "pandemic", "epidemic", "wellbeing", "medical", "doctor", "nurse", "clinic",
         "public health", "sanitation", "hygiene", "reproductive health", "disability", "HIV", "malaria",
         "tuberculosis", "maternal health", "child mortality", "universal health coverage",
         "primary healthcare", "disease prevention", "healthcare technology", "epidemiology",
         "telemedicine", "community health", "child health", "wellness", "occupational health",
         "digital health", "health monitoring", "medical instrumentation"],
    4:  ["education", "school", "learning", "teacher", "student", "literacy", "curriculum",
         "university", "vocational", "training", "scholarship", "classroom", "dropout", "enrollment",
         "early childhood", "higher education", "skills", "knowledge", "inclusive education",
         "lifelong learning", "digital literacy", "pedagogy", "e learning", "curriculum design",
         "teacher training", "stem education", "ict in education", "literacy campaigns",
         "global learning", "academic integrity", "assessment evaluation", "knowledge sharing"],
    5:  ["gender", "women", "girls", "equality", "empowerment", "discrimination", "violence against women",
         "sexual harassment", "feminism", "gender gap", "maternity", "reproductive rights",
         "female leadership", "gender based violence", "trafficking", "child marriage",
         "women empowerment", "gender mainstreaming", "equal pay", "women in stem", "gender justice",
         "inclusive policies", "gender identity", "women leadership", "equal opportunities",
         "workplace equality", "intersectionality", "representation", "reservation policy"],
    6:  ["water", "sanitation", "clean water", "drinking water", "wastewater", "sewage", "hygiene",
         "toilet", "water supply", "groundwater", "water quality", "water scarcity", "river",
         "watershed", "aquifer", "water treatment", "rainfall", "drought", "flood management",
         "sanitation systems", "water conservation", "irrigation systems", "safe drinking water",
         "groundwater management", "desalination", "water purification", "water reuse",
         "rainwater harvesting", "urban water supply", "water governance", "river basin management",
         "food water energy nexus"],
    7:  ["energy", "renewable", "solar", "wind", "electricity", "fossil fuel", "carbon neutral",
         "clean energy", "power grid", "hydropower", "biomass", "geothermal", "energy access",
         "electrification", "battery", "photovoltaic", "off-grid", "energy efficiency", "coal", "gas",
         "solar power", "wind energy", "smart grids", "clean cooking", "energy storage", "bioenergy",
         "nuclear safety", "energy innovation", "low carbon energy", "sustainable fuels",
         "energy transition", "rural electrification", "e mobility", "energy policy", "energy economics",
         "sustainable energy systems", "hydrogen energy", "transmission of energy", "smart meters",
         "fuel cell", "fuel cells", "hydrogen fuel", "electrolytic process", "thermal decomposition",
         "photochemical method", "photo catalytic method", "hydrogen storage", "hydrogen transportation",
         "hydrogen safety", "otec", "ocean thermal energy", "geothermal energy", "binary cycle power plant",
         "hydrogen sensing", "clean development mechanism", "low power systems"],
    8:  ["employment", "jobs", "economic growth", "gdp", "labor", "wages", "workforce", "decent work",
         "unemployment", "entrepreneurship", "business", "trade", "productivity", "fair wages",
         "child labor", "forced labor", "worker rights", "green economy", "tourism", "financial services",
         "decent jobs", "start ups", "green jobs", "human capital development", "corporate social responsibility",
         "workplace safety", "inclusive economy", "digital economy", "sustainable enterprises", "msme"],
    9:  ["infrastructure", "innovation", "industry", "manufacturing", "technology", "research",
         "development", "digitalization", "internet", "transport", "roads", "bridges", "railway",
         "industrialization", "startup", "patent", "automation", "artificial intelligence", "5g",
         "smart infrastructure", "research and development", "innovation ecosystems", "industry 4.0",
         "internet of things", "nanotechnology", "digital transformation", "sustainable logistics",
         "smart manufacturing", "materials science", "industrial policy", "public private partnerships",
         "sustainable engineering", "full stack development", "cyber security", "machine learning",
         "data analytics", "data mining", "digital twin technology", "drone technologies"],
    10: ["inequality", "income gap", "discrimination", "marginalized", "inclusion", "social mobility",
         "wealth distribution", "refugee", "migrant", "race", "ethnicity", "disability rights",
         "affirmative action", "equal opportunity", "progressive taxation", "remittances",
         "accessibility", "economic fairness", "anti discrimination", "migration rights",
         "representation", "fair trade", "policy advocacy", "social equity", "anti racism",
         "age inclusivity", "inclusive governance", "equity based policy", "slum upgradation"],
    11: ["urban", "city", "housing", "slum", "transport", "smart city", "public space", "sustainable city",
         "urban planning", "affordable housing", "resilience", "disaster risk", "cultural heritage",
         "air quality", "noise pollution", "green space", "public transit", "waste management",
         "smart cities", "green buildings", "public transport", "urban governance", "sustainable mobility",
         "slum upgrading", "walkable cities", "urban ecology", "land use planning", "sustainable architecture",
         "urban resilience", "inclusive urban spaces", "green infrastructure", "disaster management",
         "housing policy", "climate resilient cities", "energy efficient buildings", "sustainable urbanization"],
    12: ["consumption", "production", "waste", "recycling", "sustainable", "circular economy",
         "supply chain", "lifecycle", "eco-friendly", "plastic", "packaging", "food loss", "fast fashion",
         "chemical", "toxic", "hazardous waste", "sustainable procurement", "green product",
         "sustainable production", "resource efficiency", "waste reduction", "green design",
         "sustainable supply chains", "life cycle assessment", "responsible consumption",
         "recycling systems", "industrial ecology", "zero waste policy", "extended producer responsibility",
         "food waste reduction", "sustainable packaging", "waste to energy", "waste minimization"],
    13: ["climate", "global warming", "greenhouse gas", "carbon", "emissions", "temperature rise",
         "climate change", "adaptation", "mitigation", "paris agreement", "decarbonization",
         "carbon footprint", "net zero", "ipcc", "extreme weather", "sea level rise", "flood", "wildfire",
         "climate resilience", "mitigation strategies", "disaster risk reduction", "carbon neutrality",
         "renewable transition", "sustainable transport", "reforestation", "climate justice",
         "resilient infrastructure", "climate modelling", "low carbon economy", "climate education"],
    14: ["ocean", "marine", "sea", "fisheries", "coral reef", "plastic pollution", "overfishing",
         "aquatic", "coastal", "deep sea", "maritime", "blue economy", "seabed", "fish stock",
         "marine biodiversity", "ocean acidification", "illegal fishing", "mangrove",
         "marine ecosystems", "oceanography", "sustainable fisheries", "coastal management",
         "aquaculture", "marine pollution", "ocean governance", "sustainable shipping",
         "marine renewable energy", "marine protected areas", "marine biotechnology"],
    15: ["forest", "biodiversity", "land", "deforestation", "ecosystem", "wildlife", "species",
         "habitat", "endangered", "conservation", "soil", "desertification", "wetland", "reforestation",
         "poaching", "invasive species", "terrestrial", "national park", "land degradation",
         "forest conservation", "soil management", "desertification control", "ecosystem services",
         "afforestation", "wildlife protection", "protected areas", "sustainable forestry",
         "habitat restoration", "sustainable land management", "wetland conservation",
         "ecological monitoring", "forest policy"],
    16: ["peace", "justice", "institution", "governance", "corruption", "rule of law", "human rights",
         "democracy", "transparency", "accountability", "conflict", "violence", "crime", "legal",
         "court", "police", "refugee", "asylum", "freedom of speech", "civil society", "access to justice",
         "good governance", "justice systems", "anti corruption", "conflict resolution", "democratic institutions",
         "global security", "legal studies", "peacebuilding", "international law", "public administration",
         "civic engagement", "institutional reforms", "law enforcement", "strong institutions"],
    17: ["partnership", "cooperation", "global", "international", "financing", "technology transfer",
         "aid", "development assistance", "capacity building", "trade", "multilateral", "united nations",
         "sustainable development", "sdg", "official development assistance", "south-south",
         "global partnerships", "multi stakeholder collaboration", "international cooperation",
         "knowledge sharing", "public private partnerships", "global citizenship", "resource mobilization",
         "development finance", "research collaboration", "policy networks", "ict for development",
         "data for sdgs", "global initiatives", "open access knowledge", "sustainable financing",
         "data warehousing"],
}

SDG_CONTEXT_TERMS = {
    7: ["energy", "renewable", "solar", "wind", "power", "grid", "hydrogen", "fuel cell",
        "electrification", "battery", "geothermal", "hydropower", "smart meter", "otec"],
    11: ["urban", "city", "housing", "public transport", "smart city", "waste management",
         "slum", "walkable", "urban planning", "housing policy", "green infrastructure", "mobility"],
    12: ["waste", "recycling", "packaging", "circular economy", "life cycle", "procurement"],
    13: ["climate", "carbon", "emissions", "mitigation", "adaptation", "global warming"],
}

IMAGE_EXTENSIONS = ('png', 'jpg', 'jpeg', 'bmp', 'tiff', 'webp')

VISION_CONCEPTS = [
    {"label": "solar panels", "sdgs": [7, 13], "weight": 1.0},
    {"label": "rooftop solar panels", "sdgs": [7, 11, 13], "weight": 1.0},
    {"label": "solar farm", "sdgs": [7, 13], "weight": 1.0},
    {"label": "wind turbines", "sdgs": [7, 13], "weight": 1.0},
    {"label": "battery storage system", "sdgs": [7, 9], "weight": 0.95},
    {"label": "electrical substation", "sdgs": [7, 9], "weight": 0.9},
    {"label": "hydrogen plant", "sdgs": [7, 9], "weight": 0.95},
    {"label": "fuel cell system", "sdgs": [7, 9], "weight": 0.95},
    {"label": "geothermal power plant", "sdgs": [7, 13], "weight": 0.95},
    {"label": "hydropower dam", "sdgs": [7, 9], "weight": 0.9},
    {"label": "power lines", "sdgs": [7, 9], "weight": 0.75},
    {"label": "factory", "sdgs": [8, 9, 12], "weight": 0.9},
    {"label": "construction site", "sdgs": [8, 9, 11], "weight": 0.8},
    {"label": "bridge", "sdgs": [9, 11], "weight": 0.75},
    {"label": "road infrastructure", "sdgs": [9, 11], "weight": 0.75},
    {"label": "train", "sdgs": [9, 11], "weight": 0.75},
    {"label": "urban street", "sdgs": [11], "weight": 0.9},
    {"label": "apartment buildings", "sdgs": [11], "weight": 0.9},
    {"label": "residential neighborhood", "sdgs": [11], "weight": 0.85},
    {"label": "bus stop", "sdgs": [11], "weight": 0.8},
    {"label": "city traffic", "sdgs": [11], "weight": 0.8},
    {"label": "classroom", "sdgs": [4], "weight": 1.0},
    {"label": "students studying", "sdgs": [4], "weight": 0.95},
    {"label": "teacher teaching", "sdgs": [4], "weight": 0.95},
    {"label": "hospital", "sdgs": [3], "weight": 1.0},
    {"label": "doctor with patient", "sdgs": [3], "weight": 0.95},
    {"label": "vaccination", "sdgs": [3], "weight": 0.95},
    {"label": "women leadership meeting", "sdgs": [5, 8], "weight": 0.8},
    {"label": "clean water tap", "sdgs": [6], "weight": 1.0},
    {"label": "water treatment plant", "sdgs": [6, 9], "weight": 0.85},
    {"label": "crop field", "sdgs": [2, 15], "weight": 1.0},
    {"label": "farmer harvesting", "sdgs": [2, 8], "weight": 0.9},
    {"label": "greenhouse farming", "sdgs": [2, 9], "weight": 0.8},
    {"label": "recycling bins", "sdgs": [12, 11], "weight": 1.0},
    {"label": "plastic waste", "sdgs": [12, 14], "weight": 0.95},
    {"label": "landfill", "sdgs": [11, 12], "weight": 0.85},
    {"label": "city skyline", "sdgs": [11, 9], "weight": 0.75},
    {"label": "slum housing", "sdgs": [1, 10, 11], "weight": 1.0},
    {"label": "public transport", "sdgs": [11, 9], "weight": 0.85},
    {"label": "flood disaster", "sdgs": [11, 13], "weight": 1.0},
    {"label": "wildfire", "sdgs": [13, 15], "weight": 1.0},
    {"label": "drought", "sdgs": [2, 6, 13], "weight": 0.95},
    {"label": "forest", "sdgs": [15, 13], "weight": 0.9},
    {"label": "wildlife", "sdgs": [15], "weight": 0.95},
    {"label": "ocean pollution", "sdgs": [14, 12], "weight": 1.0},
    {"label": "fishing boat", "sdgs": [14, 8], "weight": 0.8},
    {"label": "coral reef", "sdgs": [14], "weight": 0.95},
    {"label": "courtroom", "sdgs": [16], "weight": 0.95},
    {"label": "police", "sdgs": [16], "weight": 0.8},
    {"label": "community partnership", "sdgs": [17, 11], "weight": 0.9},
    {"label": "international conference", "sdgs": [17], "weight": 0.9},
]

VISION_LABELS = [concept["label"] for concept in VISION_CONCEPTS]
VISION_LABEL_MAP = {concept["label"]: concept for concept in VISION_CONCEPTS}


# ──────────────────────────────────────────────
#  Text Extraction
# ──────────────────────────────────────────────
def extract_text_from_file(file, filename):
    ext = filename.rsplit('.', 1)[-1].lower()

    if ext == 'pdf':
        return extract_pdf(file)
    elif ext == 'docx':
        return extract_docx(file)
    elif ext in ('txt', 'md'):
        return file.read().decode('utf-8', errors='ignore')
    elif ext in ('html', 'htm'):
        return extract_html(file)
    elif ext in IMAGE_EXTENSIONS:
        return extract_image_text(file)
    else:
        return file.read().decode('utf-8', errors='ignore')


def extract_pdf(file):
    text = []
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
        tmp.write(file.read())
        tmp_path = tmp.name
    try:
        with pdfplumber.open(tmp_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
    finally:
        os.unlink(tmp_path)
    return '\n'.join(text)


def extract_docx(file):
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp.write(file.read())
        tmp_path = tmp.name
    try:
        doc = docx.Document(tmp_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return '\n'.join(paragraphs)
    finally:
        os.unlink(tmp_path)


def extract_html(file):
    html = file.read().decode('utf-8', errors='ignore')
    soup = BeautifulSoup(html, 'html.parser')
    for tag in soup(['script', 'style', 'meta', 'head']):
        tag.decompose()
    return soup.get_text(separator='\n')


def extract_image_text(file):
    import numpy as np

    image = load_image(file)
    reader = get_ocr_reader()
    variants = build_ocr_variants(image)
    line_candidates = []

    for index, variant in enumerate(variants):
        read_kwargs = {'detail': 1, 'paragraph': True}
        if index > 0:
            read_kwargs.update({
                'contrast_ths': 0.05,
                'adjust_contrast': 0.7,
                'text_threshold': 0.6,
                'low_text': 0.3,
                'width_ths': 0.7,
            })

        result = reader.readtext(np.array(variant), **read_kwargs)
        collect_ocr_candidates(result, line_candidates)

        if not line_candidates:
            fallback = reader.readtext(np.array(variant), detail=1, paragraph=False)
            collect_ocr_candidates(fallback, line_candidates)

    return "\n".join(select_best_ocr_lines(line_candidates))


def build_ocr_variants(image):
    base = ImageOps.exif_transpose(image).convert('RGB')
    width, height = base.size
    scale = max(1.0, min(3.0, 1800 / max(width, height, 1)))
    resized = base.resize(
        (max(1, int(width * scale)), max(1, int(height * scale))),
        Image.Resampling.LANCZOS
    )

    grayscale = ImageOps.grayscale(resized)
    autocontrast = ImageOps.autocontrast(grayscale)
    sharpened = autocontrast.filter(ImageFilter.SHARPEN)
    thresholded = sharpened.point(lambda px: 255 if px > 165 else 0)

    return [resized, autocontrast, sharpened, thresholded]


def normalize_ocr_line(line):
    if not line:
        return ""
    line = line.replace('|', 'I').replace('0', 'O')
    line = re.sub(r'\s+', ' ', line).strip()
    if len(line) < 2:
        return ""
    return line


def collect_ocr_candidates(result, line_candidates):
    for item in result:
        if not isinstance(item, (list, tuple)) or len(item) < 3:
            continue
        _, text, confidence = item
        normalized = normalize_ocr_line(text)
        if not normalized:
            continue
        line_candidates.append((normalized, float(confidence or 0)))


def select_best_ocr_lines(line_candidates):
    if not line_candidates:
        return []

    ordered = sorted(line_candidates, key=lambda item: (item[1], len(item[0])), reverse=True)
    selected = []

    for text, confidence in ordered:
        if any(ocr_lines_similar(text, existing) for existing in selected):
            continue
        selected.append(text)
        if len(selected) >= 12:
            break

    return selected


def ocr_lines_similar(left, right):
    left_key = re.sub(r'[^a-z0-9]+', '', left.lower())
    right_key = re.sub(r'[^a-z0-9]+', '', right.lower())
    if not left_key or not right_key:
        return False
    if left_key == right_key:
        return True
    return SequenceMatcher(None, left_key, right_key).ratio() >= 0.88


def load_image(file):
    try:
        image = Image.open(file.stream)
    except AttributeError:
        image = Image.open(file)
    return ImageOps.exif_transpose(image).convert('RGB')


@lru_cache(maxsize=1)
def get_ocr_reader():
    try:
        import easyocr
    except ImportError as exc:
        raise RuntimeError(
            "OCR requires the 'easyocr' package to be installed from requirements.txt."
        ) from exc
    return easyocr.Reader(['en'], gpu=False, verbose=False)


@lru_cache(maxsize=1)
def get_image_classifier():
    model_name = os.getenv('SDG_VISION_MODEL', 'openai/clip-vit-base-patch32')
    try:
        import torch
        from transformers import pipeline
    except ImportError as exc:
        raise RuntimeError(
            "Vision analysis requires the 'torch' and 'transformers' packages to be installed."
        ) from exc
    device = 0 if torch.cuda.is_available() else -1
    return pipeline('zero-shot-image-classification', model=model_name, device=device)


def compute_visual_scores(image):
    classifier = get_image_classifier()
    predictions = classifier(image, candidate_labels=VISION_LABELS)
    if isinstance(predictions, dict):
        predictions = predictions.get('labels', [])

    visual_scores = defaultdict(float)
    concepts_by_sdg = defaultdict(list)
    detected_concepts = []

    for item in predictions:
        label = item.get('label')
        score = float(item.get('score', 0))
        concept = VISION_LABEL_MAP.get(label)
        if not concept or score < 0.12:
            continue

        detected_concepts.append({
            'label': label,
            'score': round(score * 100, 1),
            'sdgs': concept['sdgs'],
        })

        for sdg_num in concept['sdgs']:
            visual_scores[sdg_num] += score * concept['weight']
            concepts_by_sdg[sdg_num].append(label)

    max_score = max(visual_scores.values()) if visual_scores else 1
    normalized = {}
    for i in range(1, 18):
        raw = visual_scores.get(i, 0)
        normalized[i] = round((raw / max_score) * 100) if max_score > 0 else 0

    for sdg_num, labels in concepts_by_sdg.items():
        concepts_by_sdg[sdg_num] = list(dict.fromkeys(labels))

    detected_concepts.sort(key=lambda item: item['score'], reverse=True)
    return normalized, concepts_by_sdg, detected_concepts[:8]


def clean_text(text):
    text = text.replace('–', '-').replace('—', '-').replace('&', ' and ')
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s\-.,;:!?()]', ' ', text)
    return text.strip().lower()


# ──────────────────────────────────────────────
#  SDG Scoring Engine
# ──────────────────────────────────────────────
def compute_sdg_scores(text):
    cleaned = clean_text(text)
    words = cleaned.split()
    word_set = set(words)

    scores = defaultdict(int)
    matched_keywords = defaultdict(list)

    for sdg_num, keywords in SDG_KEYWORDS.items():
        for kw in keywords:
            kw_lower = kw.lower()
            # Single word
            if ' ' not in kw_lower:
                if kw_lower in word_set:
                    count = words.count(kw_lower)
                    scores[sdg_num] += count * 2
                    if kw not in matched_keywords[sdg_num]:
                        matched_keywords[sdg_num].append(kw)
            else:
                # Multi-word phrase
                count = len(re.findall(rf'(?<!\w){re.escape(kw_lower)}(?!\w)', cleaned))
                if count > 0:
                    scores[sdg_num] += count * 5
                    if kw not in matched_keywords[sdg_num]:
                        matched_keywords[sdg_num].append(kw)

    for sdg_num, terms in SDG_CONTEXT_TERMS.items():
        hits = 0
        for term in terms:
            term_clean = term.lower()
            if ' ' in term_clean:
                hits += len(re.findall(rf'(?<!\w){re.escape(term_clean)}(?!\w)', cleaned))
            else:
                hits += words.count(term_clean)
        if hits:
            scores[sdg_num] += hits * 3
            if hits >= 2:
                scores[sdg_num] += 4

    if scores[7] and scores[11]:
        energy_anchor_hits = sum(words.count(term) for term in [
            'energy', 'solar', 'wind', 'hydrogen', 'geothermal', 'battery', 'electricity'
        ])
        urban_anchor_hits = sum(words.count(term) for term in [
            'urban', 'city', 'housing', 'transport', 'mobility', 'slum'
        ])
        if energy_anchor_hits >= urban_anchor_hits + 2:
            scores[7] += 6
        elif urban_anchor_hits >= energy_anchor_hits + 2:
            scores[11] += 6

    # Normalize to 0-100
    max_score = max(scores.values()) if scores else 1
    normalized = {}
    for i in range(1, 18):
        raw = scores.get(i, 0)
        normalized[i] = round((raw / max_score) * 100) if max_score > 0 else 0

    # Collect all found keywords
    all_keywords = []
    for kws in matched_keywords.values():
        all_keywords.extend(kws)
    all_keywords = list(dict.fromkeys(all_keywords))  # deduplicate, preserve order

    return normalized, matched_keywords, all_keywords


def zero_scores():
    return {i: 0 for i in range(1, 18)}


def merge_modality_scores(text_scores, visual_scores):
    has_text_signal = any(text_scores.values())
    has_visual_signal = any(visual_scores.values())

    if has_text_signal and has_visual_signal:
        text_weight, visual_weight = 0.6, 0.4
    elif has_text_signal:
        text_weight, visual_weight = 1.0, 0.0
    elif has_visual_signal:
        text_weight, visual_weight = 0.0, 1.0
    else:
        text_weight, visual_weight = 0.0, 0.0

    combined = {
        i: round(text_scores.get(i, 0) * text_weight + visual_scores.get(i, 0) * visual_weight)
        for i in range(1, 18)
    }

    return combined, {
        'text': text_weight,
        'visual': visual_weight,
        'has_text_signal': has_text_signal,
        'has_visual_signal': has_visual_signal,
    }


def generate_summary(scores, ranked_sdgs, keywords, text_snippet, visual_concepts=None, weights=None):
    matched_sdgs = [s for s in ranked_sdgs if s['score'] > 0]
    kw_sample = keywords[:10]
    active = sum(1 for v in scores.values() if v > 0)
    visual_concepts = visual_concepts or []

    if not matched_sdgs:
        return (
            "No strong SDG alignment was detected from the current OCR and vision signals. "
            "Try a clearer image, richer document text, or expanded SDG mappings for better coverage."
        )

    top = matched_sdgs[:3]
    top_names = [SDG_INFO[s['num']]['name'] for s in top]

    summary = f"This document shows strongest alignment with {top_names[0]}"
    if len(top_names) > 1:
        summary += f", followed by {top_names[1]}"
    if len(top_names) > 2:
        summary += f" and {top_names[2]}"
    summary += ". "

    if kw_sample:
        summary += f"Key terms identified include: {', '.join(kw_sample[:6])}. "

    if visual_concepts:
        concept_names = [item['label'] for item in visual_concepts[:4]]
        summary += f"Visual concepts detected include: {', '.join(concept_names)}. "

    if weights and weights['has_text_signal'] and weights['has_visual_signal']:
        summary += (
            f"The final ranking blends OCR text ({int(weights['text'] * 100)}%) "
            f"and visual evidence ({int(weights['visual'] * 100)}%). "
        )
    elif weights and weights['has_visual_signal'] and not weights['has_text_signal']:
        summary += "The final ranking is driven by visual evidence because OCR text was limited. "
    elif weights and weights['has_text_signal'] and not weights['has_visual_signal']:
        summary += "The final ranking is driven by extracted text because no strong visual concepts were detected. "

    summary += f"The analysis detected relevance across {active} of the 17 SDGs, with {len(keywords)} unique keywords matched."
    return summary


# ──────────────────────────────────────────────
#  Routes
# ──────────────────────────────────────────────
@app.route('/')
def index():
    return render_template('index.html')


@app.route('/analyze', methods=['POST'])
def analyze():
    text = ''
    filename = 'pasted_text.txt'
    visual_scores = zero_scores()
    visual_concepts_by_sdg = defaultdict(list)
    detected_concepts = []
    visual_error = None
    ocr_error = None
    analysis_mode = 'text'

    if 'file' in request.files and request.files['file'].filename:
        f = request.files['file']
        filename = f.filename
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        if ext in IMAGE_EXTENSIONS:
            analysis_mode = 'multimodal'
            try:
                text = extract_image_text(f)
            except Exception as e:
                ocr_error = str(e)
                text = ''

            try:
                try:
                    f.stream.seek(0)
                except (AttributeError, OSError):
                    pass
                image = load_image(f)
                visual_scores, visual_concepts_by_sdg, detected_concepts = compute_visual_scores(image)
            except Exception as e:
                visual_error = str(e)
        else:
            try:
                text = extract_text_from_file(f, filename)
            except Exception as e:
                return jsonify({'error': f'Failed to extract text: {str(e)}'}), 400

    elif 'text' in request.form and request.form['text'].strip():
        text = request.form['text'].strip()

    else:
        return jsonify({'error': 'No file or text provided'}), 400

    if not text.strip() and not any(visual_scores.values()):
        return jsonify({'error': 'Could not extract any text from the file'}), 400

    # Score
    if text.strip():
        text_scores, matched_kw_by_sdg, all_keywords = compute_sdg_scores(text)
    else:
        text_scores = zero_scores()
        matched_kw_by_sdg = defaultdict(list)
        all_keywords = []

    scores, score_weights = merge_modality_scores(text_scores, visual_scores)

    # Build ranked SDG list
    ranked = sorted(
        [{'num': n, **SDG_INFO[n], 'score': scores[n],
          'keywords': matched_kw_by_sdg.get(n, []),
          'concepts': visual_concepts_by_sdg.get(n, [])} for n in range(1, 18)],
        key=lambda x: x['score'], reverse=True
    )

    summary = generate_summary(scores, ranked, all_keywords, text[:500], detected_concepts, score_weights)
    top = ranked[0]
    confidence = round(min(100, (top['score'] / 100) * 80 + 20)) if top['score'] > 0 else 0

    result = {
        'scores': scores,
        'ranked': ranked,
        'top_sdg': {
            'num': top['num'],
            'name': top['name'],
            'color': top['color'],
            'score': top['score'],
            'confidence': confidence,
        },
        'keywords': all_keywords,
        'visual_concepts': detected_concepts,
        'text_scores': text_scores,
        'visual_scores': visual_scores,
        'score_weights': score_weights,
        'total_keywords': len(all_keywords),
        'sdgs_matched': sum(1 for v in scores.values() if v > 0),
        'summary': summary,
        'text_length': len(text),
        'filename': filename,
        'analysis_mode': analysis_mode,
        'visual_error': visual_error,
        'ocr_error': ocr_error,
    }

    return jsonify(result)


@app.route('/export/txt', methods=['POST'])
def export_txt():
    data = request.get_json()
    ranked = data.get('ranked', [])
    keywords = data.get('keywords', [])
    summary = data.get('summary', '')
    filename = data.get('filename', 'document')

    lines = [
        "=" * 60,
        "SDG ANALYZER — FULL REPORT",
        "=" * 60,
        f"Document: {filename}",
        "",
        "ANALYSIS SUMMARY",
        "-" * 40,
        summary,
        "",
        "SDG SCORES (Ranked Highest to Lowest)",
        "-" * 40,
    ]
    for s in ranked:
        bar = '█' * (s['score'] // 10) + '░' * (10 - s['score'] // 10)
        lines.append(f"SDG {s['num']:02d} | {bar} {s['score']:3d} | {s['name']}")

    lines += [
        "",
        "DETECTED KEYWORDS",
        "-" * 40,
        ', '.join(keywords) if keywords else 'None detected',
        "",
        "=" * 60,
    ]

    content = '\n'.join(lines)
    buf = io.BytesIO(content.encode('utf-8'))
    return send_file(buf, mimetype='text/plain',
                     as_attachment=True, download_name='sdg-analysis.txt')


@app.route('/export/json', methods=['POST'])
def export_json():
    data = request.get_json()
    buf = io.BytesIO(json.dumps(data, indent=2).encode('utf-8'))
    return send_file(buf, mimetype='application/json',
                     as_attachment=True, download_name='sdg-analysis.json')


@app.route('/export/docx', methods=['POST'])
def export_docx():
    data = request.get_json() or {}
    ranked = data.get('ranked', [])
    keywords = data.get('keywords', [])
    summary = data.get('summary', '')
    filename = data.get('filename', 'document')
    top_sdg = data.get('top_sdg', {})
    score_weights = data.get('score_weights', {})
    visual_concepts = data.get('visual_concepts', [])
    charts = data.get('charts', {})

    doc = docx.Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SDG Analyzer Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run(f"Document: {filename}")
    subrun.italic = True
    subrun.font.size = Pt(11)
    subrun.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph("")

    def add_heading(text, level=1):
        p = doc.add_paragraph()
        p.style = doc.styles['Heading 1' if level == 1 else 'Heading 2']
        r = p.add_run(text)
        r.font.color.rgb = RGBColor(8, 145, 178) if level == 1 else RGBColor(30, 64, 175)
        return p

    add_heading("Executive Summary")
    summary_p = doc.add_paragraph(summary or "No summary available.")
    summary_p.paragraph_format.space_after = Pt(10)

    add_heading("Highlights")
    highlights = [
        f"Top SDG: SDG {top_sdg.get('num', '-')}: {top_sdg.get('name', 'N/A')}",
        f"Confidence: {top_sdg.get('confidence', 0)}%",
        f"Matched SDGs: {data.get('sdgs_matched', 0)}",
        f"Keywords matched: {data.get('total_keywords', 0)}",
        f"Blend: OCR {int(score_weights.get('text', 0) * 100)}% / Vision {int(score_weights.get('visual', 0) * 100)}%",
    ]
    for item in highlights:
        doc.add_paragraph(item, style='List Bullet')

    add_heading("SDG Ranking")
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "SDG"
    hdr[1].text = "Goal"
    hdr[2].text = "Score"
    hdr[3].text = "Signals"
    for entry in ranked:
        row = table.add_row().cells
        row[0].text = f"SDG {entry.get('num', '-')}"
        row[1].text = entry.get('name', '')
        row[2].text = str(entry.get('score', 0))
        signals = []
        if entry.get('keywords'):
            signals.append("Keywords: " + ", ".join(entry['keywords'][:4]))
        if entry.get('concepts'):
            signals.append("Concepts: " + ", ".join(entry['concepts'][:3]))
        row[3].text = " | ".join(signals) if signals else "-"

    doc.add_paragraph("")
    add_heading("Detected Keywords", level=2)
    doc.add_paragraph(", ".join(keywords) if keywords else "No keywords matched.")

    doc.add_paragraph("")
    add_heading("Visual Concepts", level=2)
    if visual_concepts:
        for concept in visual_concepts:
            doc.add_paragraph(
                f"{concept.get('label', '')} ({concept.get('score', 0)}%)",
                style='List Bullet'
            )
    else:
        doc.add_paragraph("No strong visual concepts detected.")

    chart_items = [
        ("SDG Distribution", charts.get('pie')),
        ("Top SDGs", charts.get('bar')),
    ]
    valid_charts = [(label, image) for label, image in chart_items if image]
    if valid_charts:
        doc.add_paragraph("")
        add_heading("Charts")
        for label, image_data in valid_charts:
            doc.add_paragraph(label, style='Heading 2')
            try:
                _, encoded = image_data.split(',', 1)
                image_bytes = io.BytesIO(base64.b64decode(encoded))
                doc.add_picture(image_bytes, width=Inches(6.4))
                last = doc.paragraphs[-1]
                last.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                doc.add_paragraph(f"Unable to embed chart: {label}")
            doc.add_paragraph("")

    if data.get('ocr_error') or data.get('visual_error'):
        add_heading("Notes", level=2)
        if data.get('ocr_error'):
            doc.add_paragraph(f"OCR: {data['ocr_error']}")
        if data.get('visual_error'):
            doc.add_paragraph(f"Vision: {data['visual_error']}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_name = re.sub(r'[^A-Za-z0-9._-]+', '-', os.path.splitext(filename)[0]).strip('-') or 'sdg-analysis'
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'{safe_name}-report.docx'
    )


if __name__ == '__main__':
    print("\n🌍 SDG Analyzer running at http://localhost:5000\n")
    app.run(debug=True, port=5000)
